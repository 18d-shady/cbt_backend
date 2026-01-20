#admin.py
import io
from django.contrib import messages
import re
from django.core.files.base import ContentFile
from django.urls import path, reverse
from django.shortcuts import render, redirect
from django.contrib import admin
from django.contrib.auth.models import User
from .models import (
    Course, QuestionImage, Exam, Question, 
    StudentAnswer, ExamSession, StudentClass, StudentScore, CourseRegistration
)
from django.utils.html import format_html
from django.utils.text import slugify
from django.http import HttpResponse
from docx import Document
from docx.shared import Inches
from docx.shared import Pt, RGBColor
import io
from .views import grade_essays
import openpyxl
from openpyxl.styles import Font
from django.db.models import Sum
from xhtml2pdf import pisa
from django.template.loader import get_template
from .admin_base import SchoolScopedAdmin, is_school_admin, is_superadmin
from .admin_users import SchoolAdmin, CustomUserAdmin # Triggers registration

from unfold.admin import ModelAdmin # Ensure you use this
from unfold.decorators import action
from unfold.admin import TabularInline




@admin.register(Course)
class CourseAdmin(SchoolScopedAdmin, ModelAdmin):
    list_display = ("name", "code", "target_class", "get_school_type")
    actions = ['clone_to_classes',]

    def get_school_type(self, obj):
        # If we are looking at an existing object
        if obj and obj.school:
            return obj.school.get_school_type_display()
        return "-"
    get_school_type.short_description = "School Type"

    def get_fields(self, request, obj=None):
        # Start with all fields from the model
        fields = list(super().get_fields(request, obj))
        
        # We only apply hiding logic for School Admins
        # Superadmins should still see everything to fix data if needed
        if is_school_admin(request.user) and not is_superadmin(request.user):
            school_type = request.user.userprofile.school.school_type
            
            if school_type == 'secondary':
                # Hide 'code', keep 'name' and 'student_class'
                if 'code' in fields:
                    fields.remove('code')
            
            else: # 'others'
                # Hide both extra fields, keep only 'name'
                if 'code' in fields: fields.remove('code')
                if 'target_class' in fields: fields.remove('target_class')
        
        return fields

    def get_list_display(self, request):
        # You can even hide columns in the table view!
        list_display = list(self.list_display)
        if is_school_admin(request.user) and not is_superadmin(request.user):
            school_type = request.user.userprofile.school.school_type
            if school_type == 'secondary' and 'code' in list_display:
                list_display.remove('code')
        return list_display
    

    @action(description="Clone this course to other classes")
    def clone_to_classes(self, request, queryset):
        if queryset.count() > 1:
            self.message_user(request, "Please select only one course to clone at a time.", messages.ERROR)
            return
        
        course = queryset.first()
        school = request.user.userprofile.school

        # If the user has selected classes and clicked 'Confirm'
        if 'apply' in request.POST:
            class_ids = request.POST.getlist('target_classes')
            created_count = 0
            skipped_count = 0

            for class_id in class_ids:
                target_class = StudentClass.objects.get(id=class_id)
                
                # Prevent duplicates if the course already exists in that class
                new_course, created = Course.objects.get_or_create(
                    school=school,
                    name=course.name,
                    target_class=target_class,
                    defaults={'code': course.code}
                )
                if created:
                    created_count += 1
                else:
                    skipped_count += 1

            self.message_user(request, f"Cloned '{course.name}' to {created_count} classes. (Skipped {skipped_count} existing)")
            return redirect(request.get_full_path())

        # Otherwise, show the selection page
        # Filter classes to only show those that DON'T already have this course
        existing_class_ids = Course.objects.filter(school=school, name=course.name).values_list('target_class_id', flat=True)
        available_classes = StudentClass.objects.filter(school=school).exclude(id__in=existing_class_ids)

        return render(request, 'admin/clone_course_to_classes.html', {
            'course': course,
            'available_classes': available_classes,
            'opts': self.model._meta, # Required for admin breadcrumbs
        })
    
class QuestionImageInline(TabularInline):
    model = QuestionImage
    extra = 1
     
class QuestionInline(TabularInline):
    model = Question
    fields = ('question_number', 'question_type')
    # Make question_number read-only so the sequence isn't accidentally broken
    readonly_fields = ('question_number',)
    extra = 0
    ordering = ('question_number',)
    # This prevents the admin from adding new rows manually 
    # since save_model handles the count
    can_delete = True

@admin.register(Exam)
class ExamAdmin(SchoolScopedAdmin, ModelAdmin):
    #form = ExamForm # Including the date/time fix from before
    inlines = [QuestionInline]
    list_display = ("title", "course", "academic_year","total_questions", "grading_actions")
    list_filter = ("academic_year", "course")
    
    def get_urls(self):
        urls = super().get_urls()
        custom_urls = [
            path('<int:exam_id>/generate-word-template/', self.generate_word_template, name="generate-word-template"),
            path('<int:exam_id>/import-word-questions/', self.import_word_questions, name="import-word-questions"),
            path('<int:exam_id>/grade-essays/', self.grade_essays_view, name="grade-essays"),
            path('<int:exam_id>/export-results/', self.export_results, name="export-exam-results"),
            path('<int:exam_id>/print-slips/', self.print_result_slips, name="print-result-slips"),
        ]
        return custom_urls + urls
    
    def save_model(self, request, obj, form, change):
        # First, save the Exam object itself
        super().save_model(request, obj, form, change)
        
        # Automatically generate Question skeletons based on total_questions
        if obj.total_questions:
            existing_questions = obj.questions.values_list('question_number', flat=True)
            new_questions = []
            
            for i in range(1, obj.total_questions + 1):
                if i not in existing_questions:
                    new_questions.append(
                        Question(
                            exam=obj,
                            school=obj.school,
                            question_number=i,
                            question_type='obj'
                        )
                    )
            
            existing_count = obj.questions.count()
            if existing_count > obj.total_questions:
                # Optional: Delete the excess questions if the admin reduced the number
                obj.questions.filter(question_number__gt=obj.total_questions).delete()
            if new_questions:
                Question.objects.bulk_create(new_questions)
                self.message_user(request, f"Automatically generated {len(new_questions)} question placeholders.")
                    
    def generate_word_template(self, request, exam_id):
        exam = self.get_object(request, exam_id)
        questions = exam.questions.all().order_by('question_number')
        
        doc = Document()
        doc.add_heading(f'Exam: {exam.title}', 0)
        
        # Instruction Box
        instruction = doc.add_paragraph()
        run = instruction.add_run("IMPORTANT INSTRUCTIONS:\n")
        run.bold = True
        run.font.color.rgb = RGBColor(255, 0, 0)
        instruction.add_run("1. Do not delete the [[[ % ... % ]]] tags.\n")
        instruction.add_run("2. Paste images directly between the tags for that question.\n")
        instruction.add_run("3. Objective types (obj) must use the [[A]] style below.\n")
        instruction.add_run("4. For 'obj' types, the CORRECT tag must be just the letter (e.g., A, B, C, or D).")

        for q in questions:
            doc.add_paragraph("___________________________________________________")
            p = doc.add_paragraph()
            p.add_run(f"--- Question {q.question_number} ---").bold = True
            
            # Tags with Bold styling
            q_text = doc.add_paragraph()
            q_text.add_run("[[[% Q %]]] ").bold = True
            q_text.add_run(q.question_text or "[Enter Question Here]")
            q_text.add_run(" [[[% /Q %]]]").bold = True

            type_p = doc.add_paragraph()
            type_p.add_run("[[[% TYPE %]]] ").bold = True
            type_p.add_run(q.question_type)
            type_p.add_run(" [[[% /TYPE %]]]").bold = True

            if q.question_type == 'obj':
                doc.add_paragraph("[[[% OPTIONS %]]]").bold = True
                doc.add_paragraph(f"[[A]] {q.option_a or ''}")
                doc.add_paragraph(f"[[B]] {q.option_b or ''}")
                doc.add_paragraph(f"[[C]] {q.option_c or ''}")
                doc.add_paragraph(f"[[D]] {q.option_d or ''}")
                doc.add_paragraph("[[[% /OPTIONS %]]]").bold = True
            
            ans_p = doc.add_paragraph()
            ans_p.add_run("[[[% CORRECT %]]] ").bold = True
            ans_p.add_run(q.correct_answer or "")
            ans_p.add_run(" [[[% /CORRECT %]]]").bold = True

            pts_p = doc.add_paragraph()
            pts_p.add_run("[[[% POINTS %]]] ").bold = True
            pts_p.add_run(str(q.point))
            pts_p.add_run(" [[[% /POINTS %]]]").bold = True
                
            doc.add_paragraph("[[[% END %]]]").bold = True
            # Note: doc.add_page_break() has been removed here to allow continuous scrolling

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        response = HttpResponse(buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename=edit_{slugify(exam.title)}.docx'
        return response

    def import_word_questions(self, request, exam_id):
        if request.method == "POST":
            exam = Exam.objects.get(id=exam_id)
            word_file = request.FILES.get("word_file")
            doc = Document(word_file)
            
            # 1. Map Image IDs to binary data
            images_data = {}
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    images_data[rel.rId] = rel.target_part.blob

            # 2. Extract paragraph texts and associate images per paragraph
            full_content = []
            for para in doc.paragraphs:
                found_image = None
                if 'pic:pic' in para._p.xml:
                    img_match = re.search(r'r:embed="(rId\d+)"', para._p.xml)
                    if img_match:
                        found_image = images_data.get(img_match.group(1))
                full_content.append({'text': para.text, 'image': found_image})

            full_text_blob = "\n".join([c['text'] for c in full_content])
            blocks = full_text_blob.split("[[[% END %]]]")

            def extract_tag(tag, text):
                pattern = rf"\[\[\[% {tag} %\]\]\](.*?)\[\[\[% /{tag} %\]\]\]"
                match = re.search(pattern, text, re.DOTALL)
                return match.group(1).strip() if match else ""

            def extract_option(letter, text):
                pattern = rf"\[\[{letter}\]\](.*?)(\[\[|$)"
                match = re.search(pattern, text, re.DOTALL)
                return match.group(1).strip() if match else ""

            highest_num = 0
            content_cursor = 0

            for block in blocks:
                q_match = re.search(r"--- Question (\d+) ---", block)
                if q_match:
                    q_num = int(q_match.group(1))
                    highest_num = max(highest_num, q_num)
                    
                    q_type = extract_tag("TYPE", block).lower()
                    q_text = extract_tag("Q", block)
                    opt_block = extract_tag("OPTIONS", block)
                    q_points = extract_tag("POINTS", block) or "1.0"
                    
                    # Update Question data
                    question, created = Question.objects.update_or_create(
                        exam=exam,
                        question_number=q_num,
                        defaults={
                            'school': exam.school,
                            'question_type': q_type,
                            'question_text': q_text,
                            'option_a': extract_option('A', opt_block) if q_type == 'obj' else None,
                            'option_b': extract_option('B', opt_block) if q_type == 'obj' else None,
                            'option_c': extract_option('C', opt_block) if q_type == 'obj' else None,
                            'option_d': extract_option('D', opt_block) if q_type == 'obj' else None,
                            'correct_answer': extract_tag("CORRECT", block).upper(),
                            'point': float(q_points),
                        }
                    )

                    # Handle Images: Check content between start of this question and the END tag
                    # We delete existing images first to prevent stacking copies on re-upload
                    QuestionImage.objects.filter(question=question).delete()
                    
                    for i in range(content_cursor, len(full_content)):
                        item = full_content[i]
                        if item['image']:
                            img_name = f"exam_{exam.id}_q{q_num}.png"
                            QuestionImage.objects.create(
                                question=question, 
                                image=ContentFile(item['image'], name=img_name)
                            )
                        
                        if "[[[% END %]]]" in item['text']:
                            content_cursor = i + 1
                            break
            
            if highest_num > exam.total_questions:
                exam.total_questions = highest_num
                exam.save()

            self.message_user(request, f"Processed {highest_num} questions successfully.")
            return redirect("..")
        return render(request, "admin/word_upload_form.html", {"exam_id": exam_id})
    
    def formfield_for_foreignkey(self, db_field, request, **kwargs):
        if is_school_admin(request.user) and not is_superadmin(request.user):
            school = request.user.userprofile.school
            if db_field.name == "course":
                kwargs["queryset"] = Course.objects.filter(school=school)
        return super().formfield_for_foreignkey(db_field, request, **kwargs)
    
    def print_result_slips(self, request, exam_id):
        exam = self.get_object(request, exam_id)
        scores = StudentScore.objects.filter(exam=exam).select_related('user', 'user__userprofile__student_class')
        total_possible = exam.questions.aggregate(total=Sum('point'))['total'] or 0
        
        template_path = 'admin/result_slips_pdf.html'
        context = {
            'exam': exam,
            'scores': scores,
            'total_possible': total_possible,
            'school': exam.school,
        }
        
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'filename="results_{slugify(exam.title)}.pdf"'
        
        template = get_template(template_path)
        html = template.render(context)
        
        pisa_status = pisa.CreatePDF(html, dest=response)
        if pisa_status.err:
            return HttpResponse('We had some errors <pre>' + html + '</pre>')
        return response
    
    def export_results(self, request, exam_id):
        exam = self.get_object(request, exam_id)
        scores = StudentScore.objects.filter(exam=exam).select_related('user', 'user__userprofile__student_class')
        
        # Create Workbook
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Exam Results"

        # Header Row
        headers = ["Student Name", "Username/ID", "Class", "Score", "Total Possible", "Percentage (%)"]
        ws.append(headers)
        
        # Style Header
        for cell in ws[1]:
            cell.font = Font(bold=True)

        # Calculate Total Possible Points
        total_possible = exam.questions.aggregate(total=Sum('point'))['total'] or 0

        # Data Rows
        for s in scores:
            profile = getattr(s.user, 'userprofile', None)
            class_name = str(profile.student_class) if profile else "N/A"
            percentage = (s.score / total_possible * 100) if total_possible > 0 else 0
            
            ws.append([
                s.user.get_full_name(),
                s.user.username,
                class_name,
                s.score,
                total_possible,
                round(percentage, 2)
            ])

        # Response
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename={slugify(exam.title)}_results.xlsx'
        wb.save(response)
        return response
    
    def grading_actions(self, obj):
        return format_html(
            '<div style="display: flex; gap: 6px;">'
            '<a class="button" style="background-color: #10b981; color: white; border: none; padding: 5px 10px; border-radius: 4px; font-size: x-small;" href="{}">Grade Essays</a>'
            '<a class="button" style="background-color: #3b82f6; color: white; border: none; padding: 5px 10px; border-radius: 4px; font-size: x-small;" href="{}">Get Word Template</a>'
            '<a class="button" style="background-color: #6366f1; color: white; border: none; padding: 5px 10px; border-radius: 4px; font-size: x-small;" href="{}">Import Questions</a>'
            '<a class="button" style="background-color: #f59e0b; color: white; border: none; padding: 5px 10px; border-radius: 4px; font-size: x-small;" href="{}">Export Results</a>'
            '<a class="button" style="background-color: #ef4444; color: white; border: none; padding: 5px 10px; border-radius: 4px; font-size: x-small;" href="{}">Print Result Slips</a>'
            '</div>',
            reverse('admin:grade-essays', args=[obj.pk]),
            reverse('admin:generate-word-template', args=[obj.pk]),
            reverse('admin:import-word-questions', args=[obj.pk]),
            reverse('admin:export-exam-results', args=[obj.pk]),
            reverse('admin:print-result-slips', args=[obj.pk])
        )
    grading_actions.short_description = "Exam Dashboard"

    def grade_essays_view(self, request, exam_id):
        # Import your view function
        return grade_essays(request, exam_id)


@admin.register(Question)
class QuestionAdmin(SchoolScopedAdmin, ModelAdmin):
    inlines = [QuestionImageInline]
    list_display = ("question_number", "exam", "question_type", "question_text")
    list_filter = ("exam", "question_type")
    search_fields = ("question_text",)
    ordering = ("exam", "question_number")

@admin.register(CourseRegistration)
class CourseRegistrationAdmin(SchoolScopedAdmin, ModelAdmin):
    list_display = ("user", "course", "registered_on")
    
    def formfield_for_foreignkey(self, db_field, request, **kwargs):
        if is_school_admin(request.user) and not is_superadmin(request.user):
            school = request.user.userprofile.school
            if db_field.name == "user":
                kwargs["queryset"] = User.objects.filter(userprofile__school=school)
            if db_field.name == "course":
                kwargs["queryset"] = Course.objects.filter(school=school)
        return super().formfield_for_foreignkey(db_field, request, **kwargs)

# --- Results & Sessions (Read-Only or Limited for Admins) ---

@admin.register(StudentScore)
class StudentScoreAdmin(SchoolScopedAdmin):
    list_display = ("user", "exam", "score")
    readonly_fields = ("user", "exam", "score") # Scores usually shouldn't be edited manually

@admin.register(ExamSession)
class ExamSessionAdmin(SchoolScopedAdmin, ModelAdmin):
    list_display = ("user", "exam", "start_time", "end_time")
    readonly_fields = ("user", "exam", "start_time", "end_time")

    def has_add_permission(self, request): return False
    def has_change_permission(self, request, obj=None): return False


@admin.register(StudentAnswer)
class StudentAnswerAdmin(SchoolScopedAdmin, ModelAdmin):
    list_display = ("user", "question", "answer_short", "is_correct", "points_earned")
    list_filter = ("question__exam", "is_correct")
    readonly_fields = [f.name for f in StudentAnswer._meta.fields] # Make EVERYTHING read-only

    def answer_short(self, obj):
        return obj.answer_text[:50] + "..." if len(obj.answer_text) > 50 else obj.answer_text

    def has_add_permission(self, request): return False
    def has_change_permission(self, request, obj=None): return False
    def has_delete_permission(self, request, obj=None): return False


# Re-register User
admin.site.unregister(User)
admin.site.register(User, CustomUserAdmin)
